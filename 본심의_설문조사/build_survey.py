# -*- coding: utf-8 -*-
"""
휴머노이드 연구 인프라 수요조사 설문 — DOCX 생성 스크립트
한/글 2024에서 열어 HWPX로 저장하면 됨.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_font(run, name='맑은 고딕', size=10, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_para(doc, text, size=10, bold=False, align=None, space_after=2):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_heading(doc, text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def set_cell(cell, text, size=9, bold=False, align_center=False):
    cell.text = ''
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 여러 줄
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            run = p.add_run(line)
        else:
            p2 = cell.add_paragraph()
            if align_center:
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p2.add_run(line)
        set_font(run, size=size, bold=bold)
        run.font.name = '맑은 고딕'


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                cell.width = Cm(widths_cm[idx])


def build():
    doc = Document()

    # 여백 좁게
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('휴머노이드 연구 인프라 수요조사 설문')
    set_font(run, size=16, bold=True)
    p.paragraph_format.space_after = Pt(6)

    # 주관기관 안내 박스 표
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = 'Table Grid'
    info = [
        ('사업명', '산업혁신기반구축사업 (산업통상자원부 / 한국산업기술진흥원)'),
        ('과제명', 'AI 휴머노이드 로봇 기술혁신 센터 구축'),
        ('주관기관', '서울대학교 (연구책임자: 김영오 교수)'),
        ('연구개발기간', '2025. 09. 01 ~ 2029. 12. 31'),
    ]
    for i, (k, v) in enumerate(info):
        set_cell(tbl.rows[i].cells[0], k, size=9, bold=True, align_center=True)
        set_cell(tbl.rows[i].cells[1], v, size=9)
    set_col_widths(tbl, [3.0, 13.5])

    # 설문 안내
    add_para(doc, '', size=9)
    add_para(doc,
             '안녕하십니까. 서울대학교 AI 공학연구원 로봇 기술혁신센터에서 구축 예정인 '
             '연구 인프라(장비)의 실제 수요 파악을 위해 설문을 진행합니다. '
             '응답 내용은 장비 도입 계획 수립 및 센터 운영 방향 설정에 소중히 활용됩니다.',
             size=9)

    add_para(doc, '▪ 소요 시간: 약 5분', size=9)
    add_para(doc, '▪ 회신 기한: ', size=9)
    add_para(doc, '▪ 회신 방법: 본 문서 작성 후 아래 이메일로 회신', size=9)
    add_para(doc, '▪ 담당자: 서정연 연구원 (yeon000905@snu.ac.kr / 010-5913-8578)', size=9)
    add_para(doc,
             '※ 본 설문에서 수집된 정보는 수요조사 목적 외에는 사용되지 않으며, '
             '관련 법령에 따라 안전하게 관리됩니다.',
             size=8.5)

    # 섹션 1. 기본 정보
    add_heading(doc, '1. 기본 정보')
    info_tbl = doc.add_table(rows=6, cols=2)
    info_tbl.style = 'Table Grid'
    rows1 = [
        ('기업/기관명', ''),
        ('소속 부서 / 직무', ''),
        ('담당자 성명', ''),
        ('담당자 이메일 / 연락처', '                    /'),
        ('응답일자', '       년     월     일'),
        ('주요 사업 분야', '☐ 제조    ☐ 물류    ☐ 서비스/홈    ☐ 로봇 플랫폼    ☐ 기타 (          )'),
    ]
    for i, (k, v) in enumerate(rows1):
        set_cell(info_tbl.rows[i].cells[0], k, size=9, bold=True, align_center=True)
        set_cell(info_tbl.rows[i].cells[1], v, size=9)
    set_col_widths(info_tbl, [4.5, 12.0])

    # 섹션 2. 관심 장비 및 사용 의지
    add_heading(doc, '2. 관심 장비 및 사용 의지')
    add_para(doc,
             '※ 실제 사용할 의향이 있는 장비만 왼쪽 「관심」 칸을 체크하고, '
             '체크한 장비에 한해 「사용 의지」와 「예상 사용 빈도」를 표기해 주세요.',
             size=8.5)

    def build_equip_table(equipments):
        t = doc.add_table(rows=len(equipments) + 1, cols=4)
        t.style = 'Table Grid'
        headers = ['관심', '장비 · 설명', '사용 의지', '예상 사용 빈도']
        for i, h in enumerate(headers):
            set_cell(t.rows[0].cells[i], h, size=9, bold=True, align_center=True)
        for i, (name, desc) in enumerate(equipments):
            set_cell(t.rows[i + 1].cells[0], '☐', size=11, align_center=True)

            # 장비·설명 셀: 장비명(10.5pt bold) + 설명(8.5pt regular)로 차별화
            cell = t.rows[i + 1].cells[1]
            cell.text = ''
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p1 = cell.paragraphs[0]
            run1 = p1.add_run(name)
            set_font(run1, size=10.5, bold=True)
            p1.paragraph_format.space_after = Pt(1)
            p1.paragraph_format.space_before = Pt(0)
            p2 = cell.add_paragraph()
            run2 = p2.add_run(desc)
            set_font(run2, size=8.5, bold=False)
            run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p2.paragraph_format.space_after = Pt(0)
            p2.paragraph_format.space_before = Pt(0)

            set_cell(t.rows[i + 1].cells[2],
                     '☐ 적극 사용\n☐ 가능성 있음\n☐ 관심 수준', size=9)
            set_cell(t.rows[i + 1].cells[3],
                     '☐ 주 1회 이상\n☐ 월 1~2회\n☐ 필요 시\n☐ 미정', size=9)
        set_col_widths(t, [1.2, 7.5, 3.8, 4.0])
        return t

    add_para(doc, '① 2026년 도입 예정', size=10, bold=True, space_after=2)
    eq_2026 = [
        ('험지 주행 테스트용 휴머노이드 로봇',
         '휴머노이드 험지 주행 정책 개발 및 실증'),
        ('보행 보조용 크레인',
         '로봇 거상 및 자세 제어 보조 장치'),
        ('원거리 모션 캡처 시스템',
         '휴머노이드 변위 측정 및 인간 동작 데이터 구축'),
    ]
    build_equip_table(eq_2026)

    add_para(doc, '', size=6)
    add_para(doc, '② 2027년 도입 예정', size=10, bold=True, space_after=2)
    eq_2027 = [
        ('휴머노이드 동작 분석기',
         '동작 정확도·반복성 평가를 위한 정밀 3차원 계측'),
        ('다이나모 시스템',
         '구동부 출력·토크·효율 등 성능 파라미터 평가'),
        ('관성 측정 장치 기반 동작 분석기',
         'IMU 웨어러블 기반, 외부 마커·카메라 없이 3차원 동작 측정'),
        ('초고속 카메라',
         'CMOS 기반 초고속 연속 촬영, 빠른 동작 프레임 단위 기록'),
        ('고성능 휴머노이드 동작 학습용 서버 클러스터',
         '반자율 프롬프트 기반 학습·생성 모델 최적화용 연산 서버'),
    ]
    build_equip_table(eq_2027)

    # 섹션 3. 활용 방안
    add_heading(doc, '3. 활용 방안 (구체 내용)')
    add_para(doc,
             '■ 적용 대상 작업:  ☐ 조작(manipulation)   ☐ 이동/보행   ☐ 물류 작업   '
             '☐ HRI   ☐ 기타 (          )', size=9)
    add_para(doc,
             '■ 환경:  ☐ 공장   ☐ 물류센터   ☐ 실내 서비스   ☐ 실외   '
             '☐ 기타 (          )', size=9)
    add_para(doc,
             '■ 구체 활용 방안 (3~5줄, 예: 데이터 수집 / 모델 학습 / 성능 검증 / 고객 PoC / 내부 연구)',
             size=9)
    # 기입란 표 (빈 셀)
    box = doc.add_table(rows=1, cols=1)
    box.style = 'Table Grid'
    set_cell(box.rows[0].cells[0], '\n\n\n', size=9)
    set_col_widths(box, [16.5])

    # 섹션 4. 서버/데이터
    add_heading(doc, '4. 서버 / 데이터')
    add_para(doc,
             '※ 섹션 2에서 「고성능 서버 클러스터」를 선택하지 않으신 경우 본 항목은 건너뛰셔도 됩니다.',
             size=8.5)
    add_para(doc,
             '■ 서버 활용 목적:  ☐ 모델 학습(pretraining/SFT)   ☐ inference/검증   '
             '☐ 데이터 처리   ☐ 사용 계획 없음', size=9)
    add_para(doc,
             '■ 필요한 데이터 유형:  ☐ 행동(trajectory)   ☐ 비디오   '
             '☐ 멀티모달(vision+language+action)   ☐ 센서(force/IMU 등)', size=9)

    # 섹션 5. 협력·지원
    add_heading(doc, '5. 협력 · 지원 의사')
    add_para(doc,
             '■ 공동 데이터 구축 참여:  ☐ 있음   ☐ 조건부 있음   ☐ 없음', size=9)
    add_para(doc,
             '■ 데이터 공유 가능 수준:  ☐ 내부만 사용   ☐ 제한적 공유   ☐ 공개 가능', size=9)
    add_para(doc, '■ 가장 필요한 지원 (우선순위 1~3위 번호 기재)', size=9)

    support_tbl = doc.add_table(rows=2, cols=6)
    support_tbl.style = 'Table Grid'
    support_cols = ['데이터 수집\n인프라', '학습용\nGPU', '평가/검증\n장비',
                    '테스트베드\n공간', '로봇\n하드웨어', '알고리즘/\n모델 지원']
    for i, h in enumerate(support_cols):
        set_cell(support_tbl.rows[0].cells[i], h, size=9, bold=True, align_center=True)
    for i in range(6):
        set_cell(support_tbl.rows[1].cells[i], '', size=11, align_center=True)
    set_col_widths(support_tbl, [2.75, 2.75, 2.75, 2.75, 2.75, 2.75])

    # 섹션 6. 추가 의견
    add_heading(doc, '6. 추가 의견 (자유 기술)')
    add_para(doc, '■ 필요한 장비/기능', size=9)
    box2 = doc.add_table(rows=1, cols=1)
    box2.style = 'Table Grid'
    set_cell(box2.rows[0].cells[0], '\n\n', size=9)
    set_col_widths(box2, [16.5])

    add_para(doc, '', size=3)
    add_para(doc, '■ 현재 어려운 점', size=9)
    box3 = doc.add_table(rows=1, cols=1)
    box3.style = 'Table Grid'
    set_cell(box3.rows[0].cells[0], '\n\n', size=9)
    set_col_widths(box3, [16.5])

    add_para(doc, '', size=4)
    add_para(doc,
             '응답해 주셔서 감사합니다.   |   회신처: yeon000905@snu.ac.kr',
             size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    out = os.path.join(os.path.dirname(__file__), '수요조사_설문지_v2.docx')
    doc.save(out)
    print(f'생성 완료: {out}')


if __name__ == '__main__':
    build()
